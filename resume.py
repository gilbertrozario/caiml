from docx import Document
import re
import json

document = Document('resume10.docx')
name=""
for para in document.paragraphs:
    my_text = para.text.strip()
    if "resume" != my_text.lower() and my_text != "":
        name=my_text
        break
print("name :"+name)
email=""
for para in document.paragraphs:
    my_text = para.text
    #print(my_text)
    matchEmail = re.search(r'[\w\.-]+@[\w\.-]+', my_text)

    if matchEmail:
        email = matchEmail.group(0)
        print("email id:" + matchEmail.group(0))
        break
phone=""
for para in document.paragraphs:
    my_text = para.text
    #print(my_text)
    phonePattern = re.search(r'\+[0-9]{10}|\+[0-9]{12}|[0-9]{10}|[0-9]{12}|\+[0-9]{2}-[0-9]{10}$', my_text)
    if phonePattern:
        phone = phonePattern.group(0)
        print("phone no:" + phonePattern.group(0))
        break
experience = {}
i=1
for para in document.paragraphs:
    my_text = para.text
    exp_year= re.search(r'(.*years.*)',my_text)
    if exp_year:
        var_address=exp_year.group(0)
        experience[i] = var_address
        print(experience)
        i=i+1
check=0
professional ={}
i=0
for para in document.paragraphs:
    my_text = para.text
    print(my_text)

    if(check==0):
        for j in my_text.split():
            if "professional" == j.lower():
                check=1
                #print(para.text)

    elif (check==1):
        if (para.text == ""):
            if len(professional)!=0:
                break
        else:
            prof_string = para.text
            professional[i]=prof_string
            i=i+1
            #print(para.text)
check=0
technical ={}
for para in document.paragraphs:
    my_text = para.text
    print(my_text)

    if(check==0):
        for j in my_text.split():
            if "technical" == j.lower() or "skill"  == j.lower():
                check=1
                #print(para.text)

    elif (check==1):
        if (para.text == ""):
            if len(technical)!=0:
                break
        else:
            techn_string = para.text
            try:
                technical_list = techn_string.split(':')
                technical[technical_list[0].strip()] = technical_list[1].strip()
            except:
                technical_list = techn_string.split()
                technical[technical_list[0].strip()] = technical_list[1].strip()
            #print(para.text)
check=0
education ={}
i=1
for para in document.paragraphs:
    my_text = para.text
    print(my_text)

    if(check==0):
        for j in my_text.split():
            if "academic" == j.lower() or "educational"  == j.lower()or "education"  == j.lower():
                check=1
                #print(para.text)Education

    elif (check==1):
        if (para.text == ""):
            if len(education)!=0:
                break
        else:
            edu_string = para.text
            education[i] = edu_string
            i=i+1



check=0
print(technical)
personal ={}
for para in document.paragraphs:
    my_text = para.text
    print(my_text)
    if (check == 0):
        for j in my_text.split():
            if "personal" == j.lower():
                check = 1
    elif (check==1):
        if (para.text == ""):
            if len(personal)!=0:
                break
        else:
            person_string = para.text
            try:
                personal_list = person_string.split(':')
                personal[personal_list[0].strip()] = personal_list[1].strip()
            except:
                personal_list = person_string.split()
                technical[personal_list[0].strip()] = personal_list[1].strip()
        print(para.text)



resumeDict =	dict(name=name, email_id=email, phone=phone,experience=experience,professional=professional, education=education, technical_skill= technical,personal_details = personal)
print(resumeDict)
with open('resume.json', 'w') as outfile:
    json.dump(resumeDict, outfile)

